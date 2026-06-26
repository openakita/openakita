"""Convert user templates with 【...】 placeholders to {{ variable }} Jinja format."""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any


PLACEHOLDER_RE = re.compile(r"【([^】]+)】")

ZH_TO_VAR: dict[str, str] = {
    "会议主题": "title",
    "会议标题": "title",
    "文档标题": "title",
    "标题": "title",
    "主题": "title",
    "会议日期": "meeting_date",
    "日期": "date",
    "YYYY/MM/DD": "date",
    "YYYY-MM-DD": "date",
    "会议时间": "meeting_time",
    "时间": "time",
    "HH:MM": "time",
    "HH:MM - HH:MM": "time_range",
    "会议地点": "location",
    "会议地点/形式": "location",
    "地点": "location",
    "主持人": "host",
    "记录人": "recorder",
    "参会人员": "participants",
    "缺席/待同步人员": "absent",
    "姓名": "person_name",
    "姓名1、姓名2、姓名3": "participants",
    "会议背景": "background",
    "会议目标": "objective",
    "会议摘要": "summary",
    "摘要": "summary",
    "概要": "summary",
    "总结": "summary",
    "核心结论": "conclusions",
    "会议结论": "conclusions",
    "二. 核心结论": "conclusions",
    "主要结论": "conclusions",
    "分主题整理": "discussion",
    "三. 分主题整理": "discussion",
    "关键问题与风险": "risk",
    "四. 关键问题与风险": "risk",
    "会后待办清单": "action_items",
    "五. 会后待办清单": "action_items",
    "下次会议关注点": "next_meeting_focus",
    "六. 下次会议关注点": "next_meeting_focus",
    "会议信息": "meeting_info",
    "会议基本信息": "meeting_info",
    "分主题": "topic_sections",
    "议题整理": "topic_sections",
    "风险与问题": "risks",
    "关键问题": "risks",
    "待办事项": "action_items",
    "会后待办": "action_items",
    "公司名称": "company_name",
    "客户": "client",
    "合作对象": "partner_name",
    "合作对象/客户": "partner_name",
    "负责人": "owner",
    "截止日期": "due_date",
    "开始日期": "start_date",
    "完成日期": "finish_date",
    "实际完成日期": "finish_date",
    "进展": "status",
    "状态": "status",
    "当前进展": "progress",
    "最新进展": "latest_update",
    "最新进展记录": "latest_update",
    "是否延期": "delayed",
    "任务描述": "task_desc",
    "任务执行人": "task_owner",
    "本次会议为什么召开，当前业务/项目处于什么阶段": "background",
    "本次会议希望确定的关键事项": "objective",
    "讨论要点": "discussion",
    "已完成/现状": "done",
    "问题与风险": "risk",
    "下一步动作": "next_step",
    "卡点/风险": "blocker",
    "公司名称、联系人、资源优势、合作背景": "partner_info",
    "下次会议时间": "next_meeting_time",
    "下次会议主题": "next_meeting_title",
    "需提前准备的资料": "next_meeting_materials",
    "待确认问题": "pending_questions",
    "风险描述": "risk_desc",
    "影响范围": "risk_impact",
    "应对方案": "risk_solution",
    "线下 / 飞书 / 腾讯会议": "location",
}

_PATTERN_RULES: list[tuple[re.Pattern[str], str]] = [
    (re.compile(r"例如[：:]"), ""),
    (re.compile(r"主题[一二三四五六七八九十\d]+[：:]?"), "topic"),
    (re.compile(r"项目/产品名称"), "project_name"),
    (re.compile(r"任务\d+"), "task"),
    (re.compile(r"市场反馈|客户需求|合作方式"), "market_info"),
    (re.compile(r"新技术|新场景|可行性"), "tech_exploration"),
    (re.compile(r"付费需求|成本|落地难度"), "feasibility"),
    (re.compile(r"验证路径"), "validation_path"),
    (re.compile(r"当前资源|推进状态"), "current_status"),
    (re.compile(r"需沟通|验证|报价"), "follow_up"),
    (re.compile(r"已有验证"), "validated"),
]


@dataclass(slots=True)
class ConvertResult:
    output_path: str
    ok: bool
    variables: list[str]
    mapping: dict[str, str] = field(default_factory=dict)
    error: str = ""
    placeholder_count: int = 0

    def to_dict(self) -> dict[str, Any]:
        return {
            "output_path": self.output_path,
            "ok": self.ok,
            "variables": self.variables,
            "mapping": self.mapping,
            "error": self.error,
            "placeholder_count": self.placeholder_count,
        }


def _normalize_hint(hint: str) -> str:
    cleaned = hint.strip()
    cleaned = re.sub(r"^[例如：:]+\s*", "", cleaned)
    cleaned = re.sub(r"\s+", "", cleaned)
    return cleaned


def _hint_to_variable(hint: str, *, counters: dict[str, int], seen: set[str]) -> str:
    normalized = _normalize_hint(hint)
    if normalized in ZH_TO_VAR:
        base = ZH_TO_VAR[normalized]
    else:
        base = ""
        for pattern, var_name in _PATTERN_RULES:
            if pattern.search(normalized):
                base = var_name
                break
        if not base:
            base = _fallback_variable_name(normalized)

    if base in seen:
        counters[base] = counters.get(base, 1) + 1
        result = f"{base}_{counters[base]}"
    else:
        seen.add(base)
        result = base

    return result


def _fallback_variable_name(hint: str) -> str:
    ascii_parts = re.findall(r"[A-Za-z_][A-Za-z0-9_]*", hint)
    if ascii_parts:
        return "_".join(ascii_parts).lower()[:40]
    cleaned = re.sub(r"[^\u4e00-\u9fff]", "", hint)[:10]
    if cleaned:
        return f"field_{abs(hash(cleaned)) % 10000:04d}"
    return "field"


def _replace_in_text(
    text: str,
    *,
    counters: dict[str, int],
    seen: set[str],
    mapping: dict[str, str],
    all_variables: set[str],
) -> str:
    def replacer(match: re.Match[str]) -> str:
        hint = match.group(1)
        var_name = _hint_to_variable(hint, counters=counters, seen=seen)
        mapping[hint] = var_name
        all_variables.add(var_name)
        return "{{ " + var_name + " }}"

    return PLACEHOLDER_RE.sub(replacer, text)


def convert_template(
    source: Path,
    output: Path,
) -> ConvertResult:
    """Convert a DOCX with 【...】 placeholders to {{ variable }} format."""
    if not source.exists():
        return ConvertResult(str(output), False, [], error="Source template not found")
    if source.suffix.lower() != ".docx":
        return ConvertResult(str(output), False, [], error="Only .docx files are supported")

    try:
        from docx import Document
    except ImportError as exc:
        return ConvertResult(str(output), False, [], error=f"python-docx required: {exc}")

    try:
        document = Document(str(source))
    except Exception as exc:
        return ConvertResult(str(output), False, [], error=f"Cannot open DOCX: {exc}")

    counters: dict[str, int] = {}
    seen: set[str] = set()
    mapping: dict[str, str] = {}
    all_variables: set[str] = set()
    placeholder_count = 0

    for paragraph in document.paragraphs:
        if PLACEHOLDER_RE.search(paragraph.text):
            for run in paragraph.runs:
                if PLACEHOLDER_RE.search(run.text):
                    original = run.text
                    placeholder_count += len(PLACEHOLDER_RE.findall(original))
                    run.text = _replace_in_text(
                        original, counters=counters, seen=seen, mapping=mapping, all_variables=all_variables
                    )

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if PLACEHOLDER_RE.search(paragraph.text):
                        for run in paragraph.runs:
                            if PLACEHOLDER_RE.search(run.text):
                                original = run.text
                                placeholder_count += len(PLACEHOLDER_RE.findall(original))
                                run.text = _replace_in_text(
                                    original, counters=counters, seen=seen, mapping=mapping, all_variables=all_variables
                                )

    if placeholder_count == 0:
        return ConvertResult(
            str(output),
            False,
            [],
            mapping=mapping,
            error="未检测到【...】格式的占位符",
            placeholder_count=0,
        )

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output))

    variables = sorted(all_variables)
    return ConvertResult(
        str(output),
        True,
        variables,
        mapping=mapping,
        placeholder_count=placeholder_count,
    )


# ---------------------------------------------------------------------------
# Reverse mapping: variable name → list of Chinese labels
# ---------------------------------------------------------------------------

VAR_TO_ZH: dict[str, list[str]] = {}
for _zh, _var in ZH_TO_VAR.items():
    if len(_zh) <= 20:
        VAR_TO_ZH.setdefault(_var, []).append(_zh)

_LABEL_VALUE_RE = re.compile(
    r"(?:^|[，,；;\n])\s*(?P<label>[^：:，,\n]{1,20})[：:]\s*(?P<value>[^\n]{1,500})",
    re.MULTILINE,
)


def find_label_value_in_text(text: str, labels: list[str]) -> str:
    """Search text for 'label：value' patterns and return the first matching value."""
    if not text or not labels:
        return ""
    for match in _LABEL_VALUE_RE.finditer(text):
        found_label = match.group("label").strip()
        for target in labels:
            if target in found_label or found_label in target:
                value = match.group("value").strip()
                if value and value not in ("无", "N/A", "-", "/"):
                    return value
    return ""


def var_to_zh_labels(var_name: str) -> list[str]:
    """Get the Chinese labels associated with a variable name."""
    return VAR_TO_ZH.get(var_name, [])
