"""Built-in DOCX starter templates per doc_type."""

from __future__ import annotations

import shutil
from collections.abc import Callable
from pathlib import Path
from typing import Any, Literal

from word_models import (
    ACCEPTANCE_REPORT_VARIABLES,
    DOC_TYPE_STARTERS,
    DOC_TYPES,
    MEETING_MINUTES_VARIABLES,
    PROPOSAL_VARIABLES,
    RESEARCH_REPORT_VARIABLES,
    SOP_VARIABLES,
    WEEKLY_REPORT_VARIABLES,
)
from word_template_engine import extract_template_vars

_PLUGIN_ROOT = Path(__file__).resolve().parent
STARTERS_DIR = _PLUGIN_ROOT / "templates" / "starters"

TemplateSource = Literal["user", "default", "none"]


def starters_dir() -> Path:
    return STARTERS_DIR


def starter_path(doc_type: str) -> Path | None:
    spec = DOC_TYPE_STARTERS.get(doc_type)
    if not spec:
        return None
    path = STARTERS_DIR / spec["file"]
    return path if path.exists() else None


def list_starter_catalog(*, plugin_id: str = "word-maker") -> list[dict[str, Any]]:
    ensure_starter_files()
    items: list[dict[str, Any]] = []
    for doc_type, spec in DOC_TYPE_STARTERS.items():
        path = STARTERS_DIR / spec["file"]
        items.append(
            {
                "doc_type": doc_type,
                "label": spec.get("zh_label") or DOC_TYPES.get(doc_type, {}).get("zh", doc_type),
                "file": spec["file"],
                "variables": list(spec.get("variables") or []),
                "available": path.exists(),
                "default_for_doc_type": bool(spec.get("default_for_doc_type")),
                "download_path": f"/api/plugins/{plugin_id}/templates/starters/{doc_type}/download",
            }
        )
    return items


def _build_starter_docx(path: Path, *, sections: list[tuple[str, str]]) -> Path:
    from docx import Document

    document = Document()
    document.add_paragraph("{{ title }}", style="Title")
    for heading, placeholder in sections:
        document.add_paragraph(heading)
        document.add_paragraph(placeholder)
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))
    return path


def build_meeting_minutes_starter(path: Path) -> Path:
    """Standard meeting minutes template (8 fields)."""
    return _build_starter_docx(
        path,
        sections=[
            ("会议信息", "{{ meeting_info }}"),
            ("一. 会议摘要", "{{ summary }}"),
            ("二. 核心结论", "{{ conclusions }}"),
            ("三. 分主题整理", "{{ topic_sections }}"),
            ("四. 关键问题与风险", "{{ risks }}"),
            ("五. 会后待办清单", "{{ action_items }}"),
            ("六. 下次会议关注点", "{{ next_meeting_focus }}"),
        ],
    )


def build_weekly_report_starter(path: Path) -> Path:
    return _build_starter_docx(
        path,
        sections=[
            ("报告周期", "{{ report_period }}"),
            ("一. 本周摘要", "{{ summary }}"),
            ("二. 关键进展", "{{ highlights }}"),
            ("三. 关键指标", "{{ metrics }}"),
            ("四. 风险与问题", "{{ risks }}"),
            ("五. 下周计划", "{{ next_week_plan }}"),
        ],
    )


def build_proposal_starter(path: Path) -> Path:
    return _build_starter_docx(
        path,
        sections=[
            ("一. 项目背景", "{{ background }}"),
            ("二. 建设目标", "{{ objective }}"),
            ("三. 方案要点", "{{ solution }}"),
            ("四. 实施计划", "{{ implementation_plan }}"),
            ("五. 预期收益", "{{ benefits }}"),
            ("六. 方案摘要", "{{ summary }}"),
        ],
    )


def build_acceptance_report_starter(path: Path) -> Path:
    return _build_starter_docx(
        path,
        sections=[
            ("客户单位", "{{ company_name }}"),
            ("项目名称", "{{ project_name }}"),
            ("一. 验收摘要", "{{ summary }}"),
            ("二. 交付成果", "{{ deliverables }}"),
            ("三. 关键指标", "{{ metrics }}"),
            ("四. 遗留问题", "{{ open_issues }}"),
            ("五. 验收结论", "{{ conclusion }}"),
        ],
    )


def build_research_report_starter(path: Path) -> Path:
    return _build_starter_docx(
        path,
        sections=[
            ("一. 调研背景", "{{ background }}"),
            ("二. 调研方法", "{{ methodology }}"),
            ("三. 调研摘要", "{{ summary }}"),
            ("四. 主要发现", "{{ findings }}"),
            ("五. 核心结论", "{{ conclusions }}"),
            ("六. 建议", "{{ recommendations }}"),
        ],
    )


def build_sop_starter(path: Path) -> Path:
    return _build_starter_docx(
        path,
        sections=[
            ("一. 目的", "{{ purpose }}"),
            ("二. 适用范围", "{{ scope }}"),
            ("三. 角色职责", "{{ roles }}"),
            ("四. 前置条件", "{{ prerequisites }}"),
            ("五. 操作步骤", "{{ procedure_steps }}"),
            ("六. 异常处理", "{{ exceptions }}"),
            ("七. 修订记录", "{{ revision_info }}"),
        ],
    )


def _write_starter_docx(path: Path, *, title: str, sections: list[tuple[str, str]]) -> None:
    from docx import Document

    document = Document()
    document.add_heading(f"{{{{ title }}}}", level=0)
    for heading, placeholder in sections:
        document.add_heading(heading, level=1)
        document.add_paragraph(placeholder)
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))


def _meeting_starter_has_hint_paragraphs(path: Path) -> bool:
    try:
        from docx import Document

        doc = Document(str(path))
        for para in doc.paragraphs:
            text = (para.text or "").strip()
            if text.startswith("建议包含") or text.startswith("用 100-200 字") or text.startswith("按议题分段"):
                return True
    except Exception:
        return True
    return False


def _starter_needs_rebuild(path: Path, expected_vars: list[str]) -> bool:
    if not path.exists():
        return True
    try:
        inspection = extract_template_vars(path)
    except Exception:
        return True
    return set(inspection.variables) != set(expected_vars)


def _meeting_starter_needs_rebuild(path: Path) -> bool:
    if _meeting_starter_has_hint_paragraphs(path):
        return True
    return _starter_needs_rebuild(path, MEETING_MINUTES_VARIABLES)


_STARTER_BUILDERS: dict[str, tuple[list[str], Callable[[Path], Path]]] = {
    "meeting_minutes": (MEETING_MINUTES_VARIABLES, build_meeting_minutes_starter),
    "weekly_report": (WEEKLY_REPORT_VARIABLES, build_weekly_report_starter),
    "proposal": (PROPOSAL_VARIABLES, build_proposal_starter),
    "acceptance_report": (ACCEPTANCE_REPORT_VARIABLES, build_acceptance_report_starter),
    "research_report": (RESEARCH_REPORT_VARIABLES, build_research_report_starter),
    "sop": (SOP_VARIABLES, build_sop_starter),
}


def ensure_starter_files() -> None:
    """Create or refresh starter DOCX files."""
    STARTERS_DIR.mkdir(parents=True, exist_ok=True)

    for doc_type, (expected_vars, builder) in _STARTER_BUILDERS.items():
        spec = DOC_TYPE_STARTERS.get(doc_type)
        if not spec:
            continue
        target = STARTERS_DIR / spec["file"]
        needs_rebuild = (
            _meeting_starter_needs_rebuild(target)
            if doc_type == "meeting_minutes"
            else _starter_needs_rebuild(target, expected_vars)
        )
        if needs_rebuild:
            try:
                builder(target)
            except Exception:
                if not target.exists():
                    raise

    simple_target = STARTERS_DIR / DOC_TYPE_STARTERS["meeting_minutes_simple"]["file"]
    if not simple_target.exists():
        try:
            _write_starter_docx(
                simple_target,
                title="会议纪要",
                sections=[
                    ("基本信息", "{{ meeting_date }}"),
                    ("参会人员", "{{ attendees }}"),
                    ("会议摘要", "{{ summary }}"),
                    ("待办事项", "{{ action_items }}"),
                ],
            )
        except Exception:
            if not simple_target.exists():
                raise

    monthly_spec = DOC_TYPE_STARTERS.get("monthly_report")
    if monthly_spec:
        monthly_target = STARTERS_DIR / monthly_spec["file"]
        if not monthly_target.exists():
            try:
                _write_starter_docx(
                    monthly_target,
                    title=monthly_spec.get("zh_label", "monthly_report"),
                    sections=[
                        ("本月摘要", "{{ summary }}"),
                        ("下月计划", "{{ next_steps }}"),
                    ],
                )
            except Exception:
                if not monthly_target.exists():
                    raise


def _ensure_starter_source(doc_type: str) -> Path:
    """Return starter path for doc_type, refreshing files best-effort."""
    spec = DOC_TYPE_STARTERS.get(doc_type)
    if not spec:
        raise ValueError(f"No starter template for doc_type: {doc_type}")
    source = STARTERS_DIR / spec["file"]
    if source.exists():
        try:
            ensure_starter_files()
        except Exception:
            pass
        return source
    try:
        ensure_starter_files()
    except Exception:
        pass
    if source.exists():
        return source
    builder = _STARTER_BUILDERS.get(doc_type, (None, None))[1]
    if builder is not None:
        try:
            builder(source)
        except Exception as exc:
            raise FileNotFoundError(f"Starter file missing: {source}") from exc
    if not source.exists():
        raise FileNotFoundError(f"Starter file missing: {source}")
    return source


def copy_starter_to_uploads(
    doc_type: str,
    uploads_dir: Path,
    *,
    name_prefix: str | None = None,
) -> tuple[Path, dict[str, Any]]:
    """Copy a starter template into workspace uploads; returns absolute path and spec."""
    source = _ensure_starter_source(doc_type)
    spec = DOC_TYPE_STARTERS[doc_type]
    uploads_dir.mkdir(parents=True, exist_ok=True)
    prefix = name_prefix or doc_type.replace("_", "-")
    dest = uploads_dir / f"{prefix}-starter.docx"
    if dest.exists():
        stem = dest.stem
        counter = 1
        while dest.exists():
            dest = uploads_dir / f"{stem}-{counter}.docx"
            counter += 1
    shutil.copy2(source, dest)
    return dest, spec


def default_starter_upload_prefix(doc_type: str) -> str:
    return doc_type.replace("_", "-") + "-starter"


def is_default_starter_upload(path: str | Path) -> bool:
    stem = Path(path).stem.lower()
    return stem.endswith("-starter") or "-starter-" in stem


def starter_upload_matches_doc_type(path: str | Path, doc_type: str) -> bool:
    stem = Path(path).stem.lower()
    prefix = default_starter_upload_prefix(doc_type).lower()
    return stem == prefix or stem.startswith(f"{prefix}-")


def resolve_template_for_project(
    doc_type: str,
    template_path: str | Path | None,
    *,
    uploads_dir: Path,
    use_default: bool = True,
) -> tuple[Path | None, TemplateSource]:
    """Resolve user template or copy system default into uploads."""
    if template_path and str(template_path).strip():
        path = Path(template_path)
        if path.exists():
            if is_default_starter_upload(path) and not starter_upload_matches_doc_type(path, doc_type):
                pass
            else:
                source: TemplateSource = "default" if is_default_starter_upload(path) else "user"
                return path.resolve(), source

    if not use_default:
        return None, "none"

    from word_models import default_starter_doc_type

    starter_key = default_starter_doc_type(doc_type)
    if not starter_key:
        return None, "none"

    dest, _spec = copy_starter_to_uploads(starter_key, uploads_dir)
    return dest.resolve(), "default"
