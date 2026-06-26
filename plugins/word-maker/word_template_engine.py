"""DOCX template inspection and rendering for word-maker."""

from __future__ import annotations

import os
import re
import shutil
import tempfile
import time
import zipfile
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Any
from xml.sax.saxutils import escape

VAR_RE = re.compile(r"{{\s*([A-Za-z_][A-Za-z0-9_.]*)\s*}}")
CONTROL_RE = re.compile(r"{%\s*(?:for|if|elif|set)\s+([^%]+?)\s*%}")


def _word_lock_file_for(path: Path) -> Path | None:
    """Return Word's lock file (~$*.docx) if present beside the output file."""
    if not path.parent.exists():
        return None
    locks = list(path.parent.glob("~$*.docx"))
    return locks[0] if locks else None


def commit_docx_bytes(output: Path, write_bytes: Any) -> Path:
    """Write DOCX via temp file; on lock/permission error use a versioned filename."""
    output = Path(output)
    output.parent.mkdir(parents=True, exist_ok=True)
    tmp = output.with_name(f".{output.stem}.tmp{output.suffix}")
    try:
        write_bytes(tmp)
        os.replace(tmp, output)
        return output
    except (PermissionError, OSError):
        if tmp.exists():
            try:
                tmp.unlink()
            except OSError:
                pass
        alt = output.with_name(f"{output.stem}-{int(time.time())}{output.suffix}")
        write_bytes(alt)
        return alt
    finally:
        if tmp.exists():
            try:
                tmp.unlink()
            except OSError:
                pass


def save_document(document: Any, output: Path) -> Path:
    return commit_docx_bytes(output, lambda p: document.save(str(p)))


def save_docxtpl(template: Any, output: Path) -> Path:
    return commit_docx_bytes(output, lambda p: template.save(str(p)))


@dataclass(slots=True)
class TemplateInspection:
    template_path: str
    variables: list[str]
    missing: list[str]
    ok: bool
    engine: str
    error: str = ""
    var_contexts: dict[str, str] | None = None

    def to_dict(self) -> dict[str, Any]:
        d = asdict(self)
        if d.get("var_contexts") is None:
            d.pop("var_contexts", None)
        return d


@dataclass(slots=True)
class RenderResult:
    output_path: str
    ok: bool
    engine: str
    missing: list[str]
    error: str = ""

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


def _read_docx_xml(path: Path) -> str:
    with zipfile.ZipFile(path, "r") as archive:
        chunks: list[str] = []
        for name in archive.namelist():
            if name.startswith("word/") and name.endswith(".xml"):
                chunks.append(archive.read(name).decode("utf-8", errors="replace"))
    return "\n".join(chunks)


def _extract_with_docxtpl(path: Path) -> set[str] | None:
    try:
        from docxtpl import DocxTemplate
    except ImportError:
        return None
    try:
        template = DocxTemplate(str(path))
        return set(template.get_undeclared_template_variables())
    except Exception:
        return None


def _extract_with_regex(path: Path) -> set[str]:
    xml = _read_docx_xml(path)
    variables = {match.group(1).split(".")[0] for match in VAR_RE.finditer(xml)}
    for match in CONTROL_RE.finditer(xml):
        fragment = match.group(1)
        for token in re.findall(r"\b[A-Za-z_][A-Za-z0-9_]*\b", fragment):
            if token not in {"for", "if", "in", "and", "or", "not", "else", "True", "False"}:
                variables.add(token)
    return variables


_HEADING_STYLE_RE = re.compile(r"^Heading\s*\d|^heading|^Title", re.IGNORECASE)
_CN_NUMBERED_SECTION_RE = re.compile(
    r"^[第]?[一二三四五六七八九十百千\d]+[、.．]\s*.+"
)


def _paragraph_is_section_heading(text: str, style_name: str) -> bool:
    if not text:
        return False
    if "{{" in text or "}}" in text:
        return False
    if _HEADING_STYLE_RE.match(style_name or ""):
        return True
    return bool(_CN_NUMBERED_SECTION_RE.match(text.strip()))


def _looks_like_section_label(text: str) -> bool:
    cleaned = text.strip()
    if not cleaned or "{{" in cleaned or "}}" in cleaned:
        return False
    if _CN_NUMBERED_SECTION_RE.match(cleaned):
        return True
    return len(cleaned) <= 24 and not VAR_RE.search(cleaned)


def build_var_descriptions(
    variables: list[str],
    var_contexts: dict[str, str] | None = None,
    *,
    extra_hints: dict[str, str] | None = None,
) -> dict[str, str]:
    """Merge template position hints, convert hints, and canonical field descriptions."""
    from word_models import TEMPLATE_VAR_HINTS
    from word_template_convert import var_to_zh_labels

    contexts = dict(var_contexts or {})
    hints = dict(extra_hints or {})
    result: dict[str, str] = {}
    for var in variables:
        parts: list[str] = []
        ctx = contexts.get(var, "").strip()
        if ctx and "{{" not in ctx:
            parts.append(ctx)
        elif hints.get(var):
            parts.append(f"模板占位提示：{hints[var]}")
        semantic = TEMPLATE_VAR_HINTS.get(var, "")
        if semantic:
            parts.append(semantic)
        elif not parts:
            labels = var_to_zh_labels(var)
            if labels:
                parts.append("常见字段：" + "、".join(labels[:3]))
        result[var] = "；".join(parts)
    return result


def extract_var_contexts(template_path: str | Path) -> dict[str, str]:
    """Extract positional context for each {{ var }} in a DOCX template.

    Returns a dict mapping variable name to a description like:
    "位于章节「二. 核心结论」下，段落提示：建议提炼 3-8 条"
    """
    path = Path(template_path)
    if not path.exists() or path.suffix.lower() != ".docx":
        return {}
    try:
        from docx import Document
    except ImportError:
        return {}

    try:
        document = Document(str(path))
    except Exception:
        return {}

    contexts: dict[str, str] = {}
    current_heading = ""
    prev_text = ""

    for paragraph in document.paragraphs:
        style_name = (paragraph.style.name or "") if paragraph.style else ""
        text = paragraph.text.strip()
        if _paragraph_is_section_heading(text, style_name):
            current_heading = text
            prev_text = text
            continue
        matches = VAR_RE.findall(text)
        if not matches:
            if text:
                prev_text = text
            continue
        section_label = current_heading
        if (not section_label or "{{" in section_label) and _looks_like_section_label(prev_text):
            section_label = prev_text
        for var_name in matches:
            var_name = var_name.split(".")[0]
            surrounding = VAR_RE.sub("", text).strip()
            parts: list[str] = []
            if section_label:
                parts.append(f"位于章节「{section_label}」下")
            if surrounding and surrounding != var_name:
                parts.append(f"段落提示：{surrounding[:120]}")
            contexts[var_name] = "，".join(parts) if parts else ""
        prev_text = text

    for table in document.tables:
        for row in table.rows:
            row_label = ""
            for ci, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                if ci == 0 and cell_text and not VAR_RE.search(cell_text):
                    row_label = cell_text
                    continue
                matches = VAR_RE.findall(cell_text)
                for var_name in matches:
                    var_name = var_name.split(".")[0]
                    if var_name in contexts:
                        continue
                    parts = []
                    if current_heading:
                        parts.append(f"位于章节「{current_heading}」下")
                    if row_label:
                        parts.append(f"表格行标题：{row_label[:60]}")
                    surrounding = VAR_RE.sub("", cell_text).strip()
                    if surrounding:
                        parts.append(f"单元格提示：{surrounding[:80]}")
                    contexts[var_name] = "，".join(parts) if parts else ""

    return contexts


def extract_template_vars(
    template_path: str | Path,
    *,
    context: dict[str, Any] | None = None,
    extra_hints: dict[str, str] | None = None,
) -> TemplateInspection:
    path = Path(template_path)
    if not path.exists():
        return TemplateInspection(str(path), [], [], False, "none", "Template not found")
    if path.suffix.lower() != ".docx":
        return TemplateInspection(str(path), [], [], False, "none", "Only DOCX templates are supported")
    try:
        variables = _extract_with_docxtpl(path)
        engine = "docxtpl" if variables is not None else "regex"
        if variables is None:
            variables = _extract_with_regex(path)
    except Exception as exc:
        return TemplateInspection(str(path), [], [], False, "docx", str(exc))

    raw_contexts = extract_var_contexts(path)
    var_contexts = build_var_descriptions(
        sorted(variables),
        raw_contexts,
        extra_hints=extra_hints,
    )

    provided = set((context or {}).keys())
    missing = sorted(var for var in variables if var not in provided)
    return TemplateInspection(
        str(path),
        sorted(variables),
        missing,
        ok=not missing,
        engine=engine,
        var_contexts=var_contexts or None,
    )


def _render_with_docxtpl(path: Path, output_path: Path, context: dict[str, Any]) -> Path | None:
    try:
        from docxtpl import DocxTemplate
    except ImportError:
        return None
    template = DocxTemplate(str(path))
    template.render(context)
    return save_docxtpl(template, Path(output_path))


def _render_xml_text(xml: str, context: dict[str, Any]) -> str:
    def replace(match: re.Match[str]) -> str:
        key = match.group(1)
        value = context
        for part in key.split("."):
            value = value.get(part, "") if isinstance(value, dict) else getattr(value, part, "")
        return escape(str(value))

    return VAR_RE.sub(replace, xml)


def _render_with_regex(path: Path, output_path: Path, context: dict[str, Any]) -> None:
    with tempfile.TemporaryDirectory() as tmp:
        tmp_dir = Path(tmp)
        with zipfile.ZipFile(path, "r") as archive:
            archive.extractall(tmp_dir)
        for xml_path in (tmp_dir / "word").glob("*.xml"):
            xml_path.write_text(
                _render_xml_text(xml_path.read_text(encoding="utf-8"), context),
                encoding="utf-8",
            )
        with zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as archive:
            for file_path in tmp_dir.rglob("*"):
                if file_path.is_file():
                    archive.write(file_path, file_path.relative_to(tmp_dir).as_posix())


def append_outline_to_template(
    template_path: str | Path,
    output_path: str | Path,
    *,
    outline: dict[str, Any],
    fields: dict[str, Any] | None = None,
) -> RenderResult:
    """When a DOCX has no {{ var }} placeholders, keep the template and append generated body."""
    try:
        from docx import Document
    except ImportError as exc:
        return RenderResult(str(output_path), False, "append", [], str(exc))

    template = Path(template_path)
    output = Path(output_path)
    if not template.exists():
        return RenderResult(str(output), False, "append", [], "Template not found")

    merged_fields = dict(fields or {})
    document = Document(str(template))
    document.add_page_break()

    title = str(merged_fields.get("title") or outline.get("title") or "").strip()
    if title:
        document.add_heading(title, level=0)

    summary_keys = ("summary", "description", "content", "body", "report_summary", "abstract")
    summary = ""
    for key in summary_keys:
        value = str(merged_fields.get(key) or "").strip()
        if value and value != title:
            summary = value
            break
    if not summary:
        from word_outline_sync import outline_summary_text

        summary = outline_summary_text(outline)
    if summary:
        document.add_paragraph(summary)

    for section in outline.get("sections") or []:
        if not isinstance(section, dict):
            continue
        section_title = str(section.get("title") or "").strip()
        if section_title:
            document.add_heading(section_title, level=1)
        goal = str(section.get("goal") or "").strip()
        if goal:
            document.add_paragraph(goal)
        for bullet in section.get("bullets") or []:
            line = str(bullet or "").strip()
            if line:
                document.add_paragraph(line, style="List Bullet")

    lock = _word_lock_file_for(output)
    if lock is not None:
        output = output.with_name(f"{output.stem}-{int(time.time())}{output.suffix}")
    actual = save_document(document, output)
    return RenderResult(str(actual), True, "append", [], "")


def render_template(
    template_path: str | Path,
    output_path: str | Path,
    context: dict[str, Any],
    *,
    allow_missing: bool = False,
    outline: dict[str, Any] | None = None,
) -> RenderResult:
    template = Path(template_path)
    output = Path(output_path)
    inspection = extract_template_vars(template, context=context)
    if not inspection.variables and outline and (outline.get("sections") or context):
        return append_outline_to_template(
            template,
            output,
            outline=outline,
            fields=context,
        )
    if not inspection.ok and not allow_missing:
        missing = ", ".join(inspection.missing)
        return RenderResult(
            str(output),
            False,
            inspection.engine,
            inspection.missing,
            f"Missing template variables: {missing}" if missing else "Missing template variables",
        )
    output.parent.mkdir(parents=True, exist_ok=True)
    lock = _word_lock_file_for(output)
    if lock is not None:
        output = output.with_name(f"{output.stem}-{int(time.time())}{output.suffix}")
    used_docxtpl = False
    try:
        docxtpl_output = _render_with_docxtpl(template, output, context)
        if docxtpl_output is not None:
            used_docxtpl = True
            output = docxtpl_output
        else:
            # Simple placeholders render without docxtpl. Complex loops still need docxtpl.
            if any("{%" in line for line in _read_docx_xml(template).splitlines()):
                shutil.copyfile(template, output)
                return RenderResult(
                    str(output),
                    False,
                    "regex",
                    inspection.missing,
                    "docxtpl is required for control-flow tags",
                )
            _render_with_regex(template, output, context)
    except (PermissionError, OSError) as exc:
        lock = _word_lock_file_for(output)
        hint = (
            f"；输出文件可能被 Word/WPS 占用（{lock.name}）"
            if lock is not None
            else ""
        )
        return RenderResult(str(output), False, inspection.engine, inspection.missing, f"{exc}{hint}")
    except Exception as exc:
        return RenderResult(str(output), False, inspection.engine, inspection.missing, str(exc))
    return RenderResult(str(output), True, "docxtpl" if used_docxtpl else "regex", inspection.missing)

